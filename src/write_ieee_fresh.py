"""
Create IEEE paper from scratch using python-docx (fresh doc, no template).
Two-column body with single-column title/abstract.
"""

from docx import Document
from docx.shared import Pt, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree
import json
import os

RESULTS_DIR = os.path.join(os.path.dirname(__file__), "..", "results")
OUTPUT = os.path.join(os.path.dirname(__file__), "..", "AI700-001-IEEE-Hallucination-LLM-Paper.docx")


def load_results():
    with open(os.path.join(RESULTS_DIR, "summary.json")) as f:
        return json.load(f)


def make_section_cols(sectPr, num_cols, space_twips=360):
    for c in sectPr.findall(qn('w:cols')):
        sectPr.remove(c)
    cols = etree.SubElement(sectPr, qn('w:cols'))
    cols.set(qn('w:num'), str(num_cols))
    cols.set(qn('w:space'), str(space_twips))


def add_section_break(doc, num_cols_this_section):
    """Add continuous section break by embedding sectPr in last paragraph."""
    last_p = doc.paragraphs[-1]._p
    pPr = last_p.find(qn('w:pPr'))
    if pPr is None:
        pPr = etree.SubElement(last_p, qn('w:pPr'))
        last_p.insert(0, pPr)

    sectPr = etree.SubElement(pPr, qn('w:sectPr'))
    # Continuous
    st = etree.SubElement(sectPr, qn('w:type'))
    st.set(qn('w:val'), 'continuous')
    # Page size
    pgSz = etree.SubElement(sectPr, qn('w:pgSz'))
    pgSz.set(qn('w:w'), '12240')
    pgSz.set(qn('w:h'), '15840')
    # Margins
    pgMar = etree.SubElement(sectPr, qn('w:pgMar'))
    pgMar.set(qn('w:top'), '1080')
    pgMar.set(qn('w:bottom'), '1440')
    pgMar.set(qn('w:left'), '900')
    pgMar.set(qn('w:right'), '900')
    pgMar.set(qn('w:header'), '720')
    pgMar.set(qn('w:footer'), '720')
    pgMar.set(qn('w:gutter'), '0')
    # Columns for THIS section
    make_section_cols(sectPr, num_cols_this_section)


def p(doc, text, size=10, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
      space_after=3, space_before=0, name='Times New Roman'):
    """Add a paragraph."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = name
    # Set east asian and complex script fonts too
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        rPr = etree.SubElement(run._element, qn('w:rPr'))
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:cs'), name)

    para.alignment = align
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.line_spacing = 1.0
    return para


def heading(doc, text, level=1):
    """Add an IEEE-style heading."""
    roman = {1: 'I', 2: 'II', 3: 'III', 4: 'IV', 5: 'V', 6: 'VI', 7: 'VII'}
    if level == 1:
        # Uppercase centered roman numeral heading
        nums = ['I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII']
        para = p(doc, text, size=10, bold=True,
                 align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=4)
    else:
        # Italic left-aligned subheading
        para = p(doc, text, size=10, bold=True, italic=True,
                 align=WD_ALIGN_PARAGRAPH.LEFT, space_before=6, space_after=3)
    return para


def mixed(doc, parts, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=3):
    """Add paragraph with mixed formatting. parts: list of (text, size, bold, italic)."""
    para = doc.add_paragraph()
    for text, size, bold, italic in parts:
        run = para.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        run.font.name = 'Times New Roman'
    para.alignment = align
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.line_spacing = 1.0
    return para


def add_results_table(doc, summary):
    """Add results table."""
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Table Grid'
    table.autofit = True

    headers = ['Method', 'Halluc. Rate', 'Factual Acc.', 'ROUGE-L']
    phases = [
        ('baseline', 'Vanilla Baseline'),
        ('cot', 'Chain-of-Thought'),
        ('few_shot', 'Few-Shot'),
        ('self_consistency', 'Self-Consistency'),
        ('rag', 'RAG Pipeline'),
        ('rag_optimized', 'RAG + Optimized'),
    ]

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.size = Pt(8)
        run.font.name = 'Times New Roman'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_idx, (key, name) in enumerate(phases, 1):
        s = summary[key]
        data = [name, f"{s['hallucination_rate']:.1%}",
                f"{s['avg_factual_accuracy']:.3f}", f"{s['avg_rouge_l']:.3f}"]
        for col_idx, val in enumerate(data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(8)
            run.font.name = 'Times New Roman'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def main():
    summary = load_results()
    baseline_acc = summary['baseline']['avg_factual_accuracy']
    best_acc = summary['rag_optimized']['avg_factual_accuracy']
    improvement = ((best_acc - baseline_acc) / baseline_acc) * 100

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(0.625)
    section.right_margin = Inches(0.625)

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)

    # ============================================================
    # SINGLE COLUMN: Title, Authors, Abstract, Keywords
    # ============================================================

    # Title
    p(doc, "Hallucination Reduction in Large Language Models:\nA Comparative Study of Prompt Engineering and Retrieval-Augmented Generation",
      size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4, space_before=0)

    # Authors
    p(doc, "Emilio Garcia Martinez",
      size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0, space_before=0)
    p(doc, "Department of Artificial Intelligence",
      size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0, space_before=0)
    p(doc, "AI700-001 | Student ID: 100783029",
      size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0, space_before=0)
    p(doc, "Professor: Reda Nacif Elalaoui",
      size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4, space_before=0)

    # Abstract
    mixed(doc, [
        ("Abstract", 10, True, True),
        ("\u2014Large Language Models (LLMs) have demonstrated remarkable capabilities in natural language understanding and generation, yet they remain prone to hallucination\u2014generating fluent but factually incorrect content. This paper presents a systematic evaluation of hallucination reduction techniques applied to LLaMA 3.2 (3B parameters) across a curated dataset of 30 TruthfulQA-inspired questions spanning five knowledge categories. We implement and compare six experimental conditions: vanilla baseline, chain-of-thought (CoT) prompting, few-shot prompting, self-consistency decoding, Retrieval-Augmented Generation (RAG), and RAG with optimized prompting. Our results demonstrate that RAG-based approaches achieve the highest factual accuracy improvement of 43.3% over baseline (0.388 to 0.556), with the RAG pipeline showing the strongest ROUGE-L scores (0.353). Self-consistency decoding provided the best hallucination rate reduction among prompt-only methods. These findings confirm that grounding LLM outputs in retrieved knowledge bases is significantly more effective than prompt engineering alone for hallucination mitigation.",
         10, False, False)
    ], space_after=3)

    # Keywords - this paragraph carries the section break
    kw = mixed(doc, [
        ("Keywords", 10, True, True),
        ("\u2014Large Language Models, Hallucination Reduction, Retrieval-Augmented Generation, Prompt Engineering, Chain-of-Thought, LLaMA, Factual Accuracy",
         10, False, True)
    ], space_after=0)

    # Embed section break in keywords paragraph
    kw_pPr = kw._p.find(qn('w:pPr'))
    if kw_pPr is None:
        kw_pPr = etree.SubElement(kw._p, qn('w:pPr'))
        kw._p.insert(0, kw_pPr)
    sectPr = etree.SubElement(kw_pPr, qn('w:sectPr'))
    st = etree.SubElement(sectPr, qn('w:type'))
    st.set(qn('w:val'), 'continuous')
    pgSz = etree.SubElement(sectPr, qn('w:pgSz'))
    pgSz.set(qn('w:w'), '12240')
    pgSz.set(qn('w:h'), '15840')
    pgMar = etree.SubElement(sectPr, qn('w:pgMar'))
    pgMar.set(qn('w:top'), '720')
    pgMar.set(qn('w:bottom'), '1440')
    pgMar.set(qn('w:left'), '900')
    pgMar.set(qn('w:right'), '900')
    pgMar.set(qn('w:header'), '0')
    pgMar.set(qn('w:footer'), '0')
    pgMar.set(qn('w:gutter'), '0')
    make_section_cols(sectPr, 1)

    # ============================================================
    # TWO COLUMN: Body
    # ============================================================

    heading(doc, "I. Introduction")

    p(doc, "Large Language Models (LLMs) such as GPT-4, LLaMA, and Claude have achieved remarkable advances in natural language processing, enabling applications in healthcare, legal analysis, education, and software development [1]. However, a critical limitation persists: the phenomenon of hallucination, defined as the generation of text that is syntactically fluent but factually incorrect or unsupported by source material [2]. In high-stakes domains, hallucinated outputs can have severe consequences\u2014from incorrect medical advice to fabricated legal citations.")

    p(doc, "Recent research has proposed several approaches to mitigate hallucinations: (1) prompt engineering techniques such as chain-of-thought reasoning and few-shot prompting [1], (2) sampling strategies like self-consistency decoding [3], and (3) retrieval-augmented generation (RAG) that grounds outputs in external knowledge bases [4]. Few studies have systematically compared these under controlled conditions on the same model.")

    p(doc, "This paper implements and evaluates six experimental conditions on LLaMA 3.2 (3B parameters) across 30 TruthfulQA-inspired questions spanning Health, Science, History, Geography, and Misconceptions. Our contributions include: (1) a reproducible experimental pipeline, (2) comparative analysis of prompt engineering versus RAG, and (3) category-level analysis revealing domain-specific patterns.")

    # II. Related Work
    heading(doc, "II. Related Work")

    p(doc, "Dang, Vu, and Nguyen [1] proposed the Prompt Sensitivity (PS) and Model Variability (MV) metrics for determining whether hallucinations originate from prompting strategies or model behavior. Their experiments with GPT-4, LLaMA 2, and DeepSeek demonstrated that chain-of-thought prompting can significantly reduce prompt-sensitive hallucinations, though some models retain high rates regardless of strategy.")

    p(doc, "Bang et al. [2] introduced HalluLens, a comprehensive hallucination benchmarking suite categorizing hallucinations as extrinsic or intrinsic. Song et al. [4] presented RAG-HAT, combining retrieval augmentation with Direct Preference Optimization (DPO) fine-tuning. Wang et al. [3] demonstrated that self-consistency decoding\u2014sampling multiple reasoning paths and selecting the most consistent\u2014improves factual accuracy over single-pass generation.")

    # III. Methodology
    heading(doc, "III. Methodology")

    heading(doc, "A. Experimental Setup", level=2)
    p(doc, "All experiments were conducted using LLaMA 3.2 (3B parameters) served locally via Ollama on an Apple M1 system with 8GB RAM. The model was run with temperature 0.7 and maximum prediction length of 300 tokens. We used a curated dataset of 30 questions inspired by TruthfulQA [5], each with verified ground-truth answers and known incorrect alternatives, distributed across five categories: Health (7), Science (8), History (5), Geography (4), and Misconceptions (6).")

    heading(doc, "B. Evaluation Metrics", level=2)
    p(doc, "We evaluated responses using three complementary metrics: (1) Hallucination Detection Rate\u2014a binary classifier combining keyword overlap analysis with ROUGE-L similarity, comparing response alignment with correct versus incorrect reference answers; (2) ROUGE-L Score\u2014measuring longest common subsequence overlap between generated responses and ground-truth answers [6]; and (3) Factual Accuracy Score\u2014a composite metric combining keyword recall (weighted 0.6) and ROUGE-L similarity (weighted 0.4), specifically designed to handle verbose responses from structured prompting methods.")

    heading(doc, "C. Phase 1: Baseline Measurement", level=2)
    p(doc, "In the baseline condition, each question was presented with a simple prompt: \"Answer the following question concisely and factually.\" This establishes vanilla LLM performance without hallucination mitigation.")

    heading(doc, "D. Phase 2: Prompt-Based Reduction", level=2)
    p(doc, "Three techniques were evaluated: (a) Chain-of-Thought (CoT) Prompting\u2014instructing step-by-step reasoning, misconception identification, and scientific evidence consultation; (b) Few-Shot Prompting\u2014providing three verified question-answer examples; (c) Self-Consistency Decoding\u2014generating three independent responses (temperature 0.8) and selecting the response with highest average ROUGE-L similarity to others.")

    heading(doc, "E. Phase 3: RAG-Based Reduction", level=2)
    p(doc, "A knowledge base of four curated reference documents was indexed using ChromaDB with the all-MiniLM-L6-v2 embedding model, yielding 27 chunks. Two conditions: (a) RAG Pipeline\u2014retrieving top 3 chunks with context-grounded prompting; (b) RAG + Optimized\u2014combining retrieval with CoT, few-shot examples, and fact-checking instructions.")

    # IV. Results
    heading(doc, "IV. Results")

    heading(doc, "A. Overall Performance", level=2)

    # Table
    p(doc, "TABLE I: Performance Comparison Across Methods",
      size=8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=2)
    add_results_table(doc, summary)
    p(doc, "", space_after=4)  # spacer after table

    p(doc, f"Table I presents results across all six conditions. The most significant finding is the progressive improvement in factual accuracy from baseline ({baseline_acc:.3f}) to RAG + Optimized Prompting ({best_acc:.3f}), representing a {improvement:.1f}% improvement. The RAG pipeline achieved the highest ROUGE-L score ({summary['rag']['avg_rouge_l']:.3f}), indicating strongest lexical overlap with ground-truth answers.")

    p(doc, f"Among prompt-only methods, few-shot prompting achieved the highest factual accuracy ({summary['few_shot']['avg_factual_accuracy']:.3f}), outperforming chain-of-thought ({summary['cot']['avg_factual_accuracy']:.3f}). Self-consistency showed improvement in hallucination rate ({summary['self_consistency']['hallucination_rate']:.1%} vs {summary['baseline']['hallucination_rate']:.1%} baseline). CoT produced very low ROUGE-L ({summary['cot']['avg_rouge_l']:.3f}) due to verbose outputs, highlighting ROUGE-L limitations for structured responses.")

    heading(doc, "B. Category-Level Analysis", level=2)
    p(doc, "Category-level analysis revealed significant variation. Geography showed the largest accuracy improvement with RAG, as factual geographic data is well-suited to retrieval-based grounding. Health maintained relatively low baseline hallucination rates, suggesting stronger pre-training knowledge. History and Misconceptions were the hardest categories, likely requiring deeper reasoning beyond simple fact retrieval.")

    heading(doc, "C. Evaluation Metric Analysis", level=2)
    p(doc, "An important finding concerns the divergence between ROUGE-L and factual accuracy metrics. CoT and RAG + Optimized produced verbose, structured responses achieving low ROUGE-L despite containing correct facts. Our composite metric, weighting keyword recall at 0.6 and ROUGE-L at 0.4, better captures verbose-but-accurate responses. This has implications for evaluation methodology: relying solely on lexical overlap may underestimate structured prompting effectiveness.")

    # V. Discussion
    heading(doc, "V. Discussion")

    p(doc, "Our results reveal key insights about hallucination reduction in smaller LLMs. First, the progressive improvement from prompt engineering to RAG confirms that grounding outputs in external knowledge is more effective than relying on parametric knowledge alone. The 43.3% factual accuracy improvement demonstrates that even 3B parameter models achieve meaningful hallucination reduction when augmented with retrieval.")

    p(doc, "Second, CoT's underperformance on smaller models contrasts with findings by Dang et al. [1] on GPT-4, suggesting capacity-dependent effectiveness\u2014smaller models may lack reasoning depth to leverage step-by-step decomposition, instead generating plausible but unfaithful reasoning chains. Third, self-consistency showed moderate improvement with minimal complexity, making it a practical first-line defense at the cost of 3x inference time.")

    p(doc, "RAG is particularly effective for factual recall questions (Geography, Health) where retrieved documents directly contain relevant information, but less effective for nuanced reasoning about misconceptions. This suggests optimal strategies should be domain-adaptive.")

    heading(doc, "A. Limitations", level=2)
    p(doc, "This study has limitations: a small dataset (30 questions), single model size (3B), purpose-built knowledge base representing an optimistic RAG scenario, and automated-only evaluation. Results may not generalize to larger LLMs where CoT may be more effective.")

    # VI. Conclusion
    heading(doc, "VI. Conclusion and Future Work")

    p(doc, f"This paper presented a systematic evaluation of hallucination reduction techniques for LLaMA 3.2 (3B parameters). Our three-phase pipeline demonstrated that RAG combined with optimized prompting achieves {improvement:.1f}% factual accuracy improvement over the vanilla baseline. Among prompt-only methods, few-shot and self-consistency showed the most consistent improvements, while CoT was less effective on this smaller model.")

    p(doc, "Future work will extend evaluation to larger models (7B, 13B, 70B parameters), implement hallucination-aware fine-tuning following the RAG-HAT approach [4], evaluate on the full TruthfulQA benchmark, and explore hybrid self-consistency + RAG approaches with NLI-based faithfulness evaluation.")

    # References
    heading(doc, "References")

    refs = [
        '[1] A.-H. Dang, T. Vu, and L.-M. Nguyen, "A survey and analysis of hallucinations in large language models," Frontiers in Artificial Intelligence, vol. 8, 2025.',
        '[2] Y. Bang et al., "HalluLens: A benchmark for LLM hallucinations," in Proc. 63rd ACL, 2025.',
        '[3] X. Wang et al., "Self-consistency improves chain of thought reasoning in language models," in Proc. ICLR, 2023.',
        '[4] J. Song et al., "RAG-HAT: A Hallucination-Aware Tuning Pipeline for LLM in Retrieval-Augmented Generation," in Proc. EMNLP, 2024.',
        '[5] S. Lin, J. Hilton, and O. Evans, "TruthfulQA: Measuring how models mimic human falsehoods," in Proc. ACL, 2022.',
        '[6] C.-Y. Lin, "ROUGE: A package for automatic evaluation of summaries," in Text Summarization Branches Out, 2004.',
        '[7] P. Lewis et al., "Retrieval-augmented generation for knowledge-intensive NLP tasks," in Proc. NeurIPS, 2020.',
        '[8] J. Wei et al., "Chain-of-thought prompting elicits reasoning in large language models," in Proc. NeurIPS, 2022.',
    ]
    for ref in refs:
        p(doc, ref, size=8, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=1)

    # Set the FINAL section (body) to 2 columns
    final_sect = doc.sections[-1]
    final_sect.page_width = Inches(8.5)
    final_sect.page_height = Inches(11)
    final_sect.top_margin = Inches(0.5)
    final_sect.bottom_margin = Inches(1.0)
    final_sect.left_margin = Inches(0.625)
    final_sect.right_margin = Inches(0.625)
    make_section_cols(final_sect._sectPr, 2, 360)

    doc.save(OUTPUT)
    print(f"IEEE paper saved to: {OUTPUT}")


if __name__ == "__main__":
    main()

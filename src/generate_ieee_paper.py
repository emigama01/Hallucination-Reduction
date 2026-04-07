"""
Generate IEEE-format paper for Hallucination Reduction in LLMs project.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
import json
import os

RESULTS_DIR = os.path.join(os.path.dirname(__file__), "..", "results")
OUTPUT_DIR = os.path.dirname(os.path.dirname(__file__))


def set_paragraph_format(paragraph, font_size=10, bold=False, italic=False,
                         alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6,
                         space_before=0, font_name='Times New Roman'):
    """Helper to format paragraphs."""
    paragraph.alignment = alignment
    pf = paragraph.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    for run in paragraph.runs:
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.italic = italic
        run.font.name = font_name


def add_formatted_paragraph(doc, text, font_size=10, bold=False, italic=False,
                            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4,
                            space_before=0):
    """Add a formatted paragraph to the document."""
    from docx.shared import Twips
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = 'Times New Roman'
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.0
    return p


def add_mixed_paragraph(doc, parts, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                        space_after=4, space_before=0):
    """Add paragraph with mixed formatting. parts is list of (text, bold, italic) tuples."""
    p = doc.add_paragraph()
    for text, bold, italic in parts:
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.bold = bold
        run.font.italic = italic
        run.font.name = 'Times New Roman'
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.0
    return p


def set_section_columns(section, num_cols=2, spacing_twips=360):
    """Set a section to have multiple columns using XML. Spacing in twips (1 inch=1440)."""
    from lxml import etree
    from docx.oxml.ns import qn

    sectPr = section._sectPr
    for cols in sectPr.findall(qn('w:cols')):
        sectPr.remove(cols)

    cols = etree.SubElement(sectPr, qn('w:cols'))
    cols.set(qn('w:num'), str(num_cols))
    cols.set(qn('w:space'), str(spacing_twips))


def insert_continuous_section_break_before(paragraph, num_cols_before=1):
    """Insert a continuous section break INSIDE the paragraph's pPr so there's no extra whitespace."""
    from lxml import etree
    from docx.oxml.ns import qn

    pPr = paragraph._p.get_or_add_pPr()
    sectPr = etree.SubElement(pPr, qn('w:sectPr'))

    sectType = etree.SubElement(sectPr, qn('w:type'))
    sectType.set(qn('w:val'), 'continuous')

    pgSz = etree.SubElement(sectPr, qn('w:pgSz'))
    pgSz.set(qn('w:w'), '12240')
    pgSz.set(qn('w:h'), '15840')

    pgMar = etree.SubElement(sectPr, qn('w:pgMar'))
    pgMar.set(qn('w:top'), '1080')
    pgMar.set(qn('w:bottom'), '1440')
    pgMar.set(qn('w:left'), '900')
    pgMar.set(qn('w:right'), '900')
    pgMar.set(qn('w:header'), '720')
    pgMar.set(qn('w:footer'), '720')
    pgMar.set(qn('w:gutter'), '0')

    cols = etree.SubElement(sectPr, qn('w:cols'))
    cols.set(qn('w:num'), str(num_cols_before))


def generate_paper():
    # Load results
    with open(os.path.join(RESULTS_DIR, "summary.json")) as f:
        summary = json.load(f)

    doc = Document()

    # Page setup - first section (title, single column)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(0.625)
    section.right_margin = Inches(0.625)

    # ---- TITLE (single column) ----
    add_formatted_paragraph(
        doc,
        "Hallucination Reduction in Large Language Models: A Comparative Study of Prompt Engineering and Retrieval-Augmented Generation",
        font_size=24, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=12, space_before=12
    )

    # ---- AUTHOR ----
    add_formatted_paragraph(
        doc,
        "Emilio Garcia Martinez",
        font_size=11, bold=False,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2
    )
    add_formatted_paragraph(
        doc,
        "Department of Artificial Intelligence",
        font_size=10, bold=False, italic=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2
    )
    add_formatted_paragraph(
        doc,
        "AI700-001 | Student ID: 100783029",
        font_size=10, bold=False,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2
    )
    add_formatted_paragraph(
        doc,
        "Professor: Reda Nacif Elalaoui",
        font_size=10, bold=False,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=12
    )

    # ---- ABSTRACT (single column) ----
    add_mixed_paragraph(doc, [
        ("Abstract", True, True),
        ("\u2014Large Language Models (LLMs) have demonstrated remarkable capabilities in natural language understanding and generation, yet they remain prone to hallucination\u2014generating fluent but factually incorrect content. This paper presents a systematic evaluation of hallucination reduction techniques applied to LLaMA 3.2 (3B parameters) across a curated dataset of 30 TruthfulQA-inspired questions spanning five knowledge categories. We implement and compare six experimental conditions: vanilla baseline, chain-of-thought (CoT) prompting, few-shot prompting, self-consistency decoding, Retrieval-Augmented Generation (RAG), and RAG with optimized prompting. Our results demonstrate that RAG-based approaches achieve the highest factual accuracy improvement of 43.3% over baseline (0.388 to 0.556), with the RAG pipeline showing the strongest ROUGE-L scores (0.353). Self-consistency decoding provided the best hallucination rate reduction among prompt-only methods. These findings confirm that grounding LLM outputs in retrieved knowledge bases is significantly more effective than prompt engineering alone for hallucination mitigation.", False, False)
    ], alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8)

    kw_para = add_mixed_paragraph(doc, [
        ("Keywords", True, True),
        ("\u2014Large Language Models, Hallucination Reduction, Retrieval-Augmented Generation, Prompt Engineering, Chain-of-Thought, LLaMA, Factual Accuracy", False, True)
    ], space_after=0)

    # Attach section break directly to the keywords paragraph (no extra space)
    insert_continuous_section_break_before(kw_para, num_cols_before=1)

    # ---- I. INTRODUCTION (now in 2-column layout) ----
    add_formatted_paragraph(doc, "I. Introduction", font_size=10, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=4)

    add_formatted_paragraph(doc,
        "Large Language Models (LLMs) such as GPT-4, LLaMA, and Claude have achieved remarkable advances in natural language processing, enabling applications in healthcare, legal analysis, education, and software development [1]. However, a critical limitation persists: the phenomenon of hallucination, defined as the generation of text that is syntactically fluent but factually incorrect or unsupported by source material [2]. In high-stakes domains, hallucinated outputs can have severe consequences—from incorrect medical advice to fabricated legal citations.")

    add_formatted_paragraph(doc,
        "Recent research has proposed several approaches to mitigate hallucinations, broadly categorized into three strategies: (1) prompt engineering techniques such as chain-of-thought reasoning and few-shot prompting [1], (2) sampling strategies like self-consistency decoding [3], and (3) retrieval-augmented generation (RAG) that grounds model outputs in external knowledge bases [4]. While each approach has shown promise individually, few studies have systematically compared their effectiveness under controlled conditions on the same model and dataset.")

    add_formatted_paragraph(doc,
        "This paper addresses this gap by implementing and evaluating six experimental conditions on LLaMA 3.2 (3B parameters), testing the effectiveness of prompt-based and RAG-based hallucination reduction techniques on a curated dataset of 30 TruthfulQA-inspired questions spanning five knowledge categories: Health, Science, History, Geography, and Misconceptions. Our contributions include: (1) a reproducible experimental pipeline for hallucination evaluation, (2) a comparative analysis of prompt engineering versus RAG-based approaches, and (3) category-level analysis revealing domain-specific effectiveness patterns.")

    # ---- II. RELATED WORK ----
    add_formatted_paragraph(doc, "II. Related Work", font_size=10, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=4)

    add_formatted_paragraph(doc,
        "Dang, Vu, and Nguyen [1] proposed the Prompt Sensitivity (PS) and Model Variability (MV) metrics for determining whether hallucinations originate from prompting strategies or model behavior. Their experiments with GPT-4, LLaMA 2, and DeepSeek demonstrated that chain-of-thought prompting can significantly reduce prompt-sensitive hallucinations, though some models retain high hallucination rates regardless of strategy, indicating the need for multi-faceted approaches.")

    add_formatted_paragraph(doc,
        "Bang et al. [2] introduced HalluLens, a comprehensive hallucination benchmarking suite that categorizes hallucinations as extrinsic (deviating from training data) or intrinsic (internal inconsistencies). Their dynamic test generation approach prevents data contamination and provides robust evaluation metrics that informed our evaluation methodology.")

    add_formatted_paragraph(doc,
        "Song et al. [4] presented RAG-HAT, a hallucination-aware tuning pipeline combining retrieval augmentation with Direct Preference Optimization (DPO) fine-tuning. Their three-stage approach—hallucination detection, resolution via GPT-4 Turbo, and preference-based fine-tuning—demonstrated significant hallucination reduction, validating the effectiveness of combining RAG with targeted training approaches.")

    add_formatted_paragraph(doc,
        "Wang et al. [3] demonstrated that self-consistency decoding—sampling multiple reasoning paths and selecting the most consistent answer—improves factual accuracy over single-pass generation, particularly for questions requiring multi-step reasoning.")

    # ---- III. METHODOLOGY ----
    add_formatted_paragraph(doc, "III. Methodology", font_size=10, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=4)

    add_formatted_paragraph(doc, "A. Experimental Setup", font_size=10, bold=True, italic=True,
                            alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=8, space_after=4)

    add_formatted_paragraph(doc,
        "All experiments were conducted using LLaMA 3.2 (3B parameters) served locally via Ollama on an Apple M1 system with 8GB RAM. The model was run with temperature 0.7 and a maximum prediction length of 300 tokens. We used a curated dataset of 30 questions inspired by TruthfulQA [5], each with verified ground-truth answers and known incorrect alternatives, distributed across five categories: Health (7), Science (8), History (5), Geography (4), and Misconceptions (6).")

    add_formatted_paragraph(doc, "B. Evaluation Metrics", font_size=10, bold=True, italic=True,
                            alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=8, space_after=4)

    add_formatted_paragraph(doc,
        "We evaluated responses using three complementary metrics: (1) Hallucination Detection Rate—a binary classifier combining keyword overlap analysis with ROUGE-L similarity, comparing response alignment with correct versus incorrect reference answers; (2) ROUGE-L Score—measuring longest common subsequence overlap between generated responses and ground-truth answers [6]; and (3) Factual Accuracy Score—a composite metric combining keyword recall (weighted 0.6) and ROUGE-L similarity (weighted 0.4) against correct and incorrect answer sets, specifically designed to handle verbose responses from structured prompting methods.")

    add_formatted_paragraph(doc, "C. Phase 1: Baseline Measurement", font_size=10, bold=True, italic=True,
                            alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=8, space_after=4)

    add_formatted_paragraph(doc,
        "In the baseline condition, each question was presented to the model with a simple prompt: \"Answer the following question concisely and factually.\" This establishes the vanilla LLM performance without any hallucination mitigation strategies.")

    add_formatted_paragraph(doc, "D. Phase 2: Prompt-Based Reduction", font_size=10, bold=True, italic=True,
                            alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=8, space_after=4)

    add_formatted_paragraph(doc,
        "Three prompt engineering techniques were evaluated: (a) Chain-of-Thought (CoT) Prompting—instructing the model to reason step-by-step, identify common misconceptions, and consult scientific evidence before providing a final answer; (b) Few-Shot Prompting—providing three verified question-answer examples demonstrating fact-checked, misconception-correcting responses; (c) Self-Consistency Decoding—generating three independent responses per question (temperature 0.8) and selecting the response with highest average ROUGE-L similarity to all other responses, following the majority consensus principle.")

    add_formatted_paragraph(doc, "E. Phase 3: RAG-Based Reduction", font_size=10, bold=True, italic=True,
                            alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=8, space_after=4)

    add_formatted_paragraph(doc,
        "A knowledge base of four curated reference documents covering health myths, scientific misconceptions, historical facts, and geography was constructed and indexed using ChromaDB with the all-MiniLM-L6-v2 embedding model. Each document was split into paragraphs, yielding 27 indexed chunks. Two RAG conditions were tested: (a) RAG Pipeline—retrieving the top 3 most relevant chunks and instructing the model to answer based solely on the provided context; (b) RAG + Optimized Prompting—combining retrieval with a structured prompt incorporating fact-checking instructions, misconception identification, chain-of-thought reasoning, and a few-shot example.")

    # ---- IV. RESULTS ----
    add_formatted_paragraph(doc, "IV. Results", font_size=10, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=4)

    add_formatted_paragraph(doc, "A. Overall Performance", font_size=10, bold=True, italic=True,
                            alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=8, space_after=4)

    # Results table
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Table Grid'

    headers = ['Method', 'Halluc. Rate', 'Factual Acc.', 'ROUGE-L']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)
                r.font.name = 'Times New Roman'

    phase_data = [
        ('baseline', 'Vanilla Baseline'),
        ('cot', 'Chain-of-Thought'),
        ('few_shot', 'Few-Shot'),
        ('self_consistency', 'Self-Consistency'),
        ('rag', 'RAG Pipeline'),
        ('rag_optimized', 'RAG + Optimized')
    ]

    for row_idx, (key, name) in enumerate(phase_data, 1):
        s = summary[key]
        table.rows[row_idx].cells[0].text = name
        table.rows[row_idx].cells[1].text = f"{s['hallucination_rate']:.1%}"
        table.rows[row_idx].cells[2].text = f"{s['avg_factual_accuracy']:.3f}"
        table.rows[row_idx].cells[3].text = f"{s['avg_rouge_l']:.3f}"
        for cell in table.rows[row_idx].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.name = 'Times New Roman'

    add_formatted_paragraph(doc, "", space_after=4)
    add_formatted_paragraph(doc,
        "TABLE I. Performance comparison across all experimental conditions. Hallucination Rate indicates the percentage of responses flagged as containing hallucinated content. Factual Accuracy is a composite score combining keyword recall and ROUGE-L similarity.",
        font_size=9, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    # Key findings
    baseline_acc = summary['baseline']['avg_factual_accuracy']
    best_acc = summary['rag_optimized']['avg_factual_accuracy']
    improvement = ((best_acc - baseline_acc) / baseline_acc) * 100

    add_formatted_paragraph(doc,
        f"Table I presents the comprehensive results across all six conditions. The most significant finding is the progressive improvement in factual accuracy from baseline ({baseline_acc:.3f}) to RAG + Optimized Prompting ({best_acc:.3f}), representing a {improvement:.1f}% improvement. The RAG pipeline achieved the highest ROUGE-L score ({summary['rag']['avg_rouge_l']:.3f}), indicating strongest lexical overlap with ground-truth answers.")

    add_formatted_paragraph(doc,
        f"Among prompt-only methods, few-shot prompting achieved the highest factual accuracy ({summary['few_shot']['avg_factual_accuracy']:.3f}), outperforming chain-of-thought ({summary['cot']['avg_factual_accuracy']:.3f}). Self-consistency decoding showed improvement over baseline in both factual accuracy ({summary['self_consistency']['avg_factual_accuracy']:.3f}) and hallucination rate ({summary['self_consistency']['hallucination_rate']:.1%} vs {summary['baseline']['hallucination_rate']:.1%}). Notably, CoT prompting produced very low ROUGE-L scores ({summary['cot']['avg_rouge_l']:.3f}) due to verbose step-by-step outputs, though its factual accuracy remained competitive, highlighting the limitation of ROUGE-L as a sole evaluation metric for structured responses.")

    add_formatted_paragraph(doc, "B. Category-Level Analysis", font_size=10, bold=True, italic=True,
                            alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=8, space_after=4)

    # Category table
    categories = ['Health', 'Science', 'History', 'Geography', 'Misconceptions']
    cat_table = doc.add_table(rows=len(categories) + 1, cols=5)
    cat_table.style = 'Table Grid'

    cat_headers = ['Category', 'N', 'Baseline Acc.', 'RAG+Opt Acc.', 'Improvement']
    for i, h in enumerate(cat_headers):
        cell = cat_table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)
                r.font.name = 'Times New Roman'

    for row_idx, cat in enumerate(categories, 1):
        base_cat = summary['baseline']['by_category'].get(cat, {})
        rag_cat = summary['rag_optimized']['by_category'].get(cat, {})
        base_acc = base_cat.get('avg_accuracy', 0)
        rag_acc = rag_cat.get('avg_accuracy', 0)
        imp = ((rag_acc - base_acc) / base_acc * 100) if base_acc > 0 else 0

        cat_table.rows[row_idx].cells[0].text = cat
        cat_table.rows[row_idx].cells[1].text = str(base_cat.get('count', 0))
        cat_table.rows[row_idx].cells[2].text = f"{base_acc:.3f}"
        cat_table.rows[row_idx].cells[3].text = f"{rag_acc:.3f}"
        cat_table.rows[row_idx].cells[4].text = f"+{imp:.1f}%"
        for cell in cat_table.rows[row_idx].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.name = 'Times New Roman'

    add_formatted_paragraph(doc, "", space_after=4)
    add_formatted_paragraph(doc,
        "TABLE II. Category-level factual accuracy comparison between baseline and RAG + Optimized Prompting conditions.",
        font_size=9, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    # Find best and worst categories
    cat_improvements = {}
    for cat in categories:
        base_acc = summary['baseline']['by_category'].get(cat, {}).get('avg_accuracy', 0)
        rag_acc = summary['rag_optimized']['by_category'].get(cat, {}).get('avg_accuracy', 0)
        cat_improvements[cat] = ((rag_acc - base_acc) / base_acc * 100) if base_acc > 0 else 0

    best_cat = max(cat_improvements, key=cat_improvements.get)
    worst_cat = min(cat_improvements, key=cat_improvements.get)

    add_formatted_paragraph(doc,
        f"Table II reveals significant variation in technique effectiveness across knowledge categories. The {best_cat} category showed the largest improvement ({cat_improvements[best_cat]:.1f}%), while {worst_cat} showed the least improvement ({cat_improvements[worst_cat]:.1f}%). Geography questions benefited most from RAG, likely because factual geographic data is well-suited to retrieval-based grounding. The Health category maintained relatively low baseline hallucination rates, suggesting the model's pre-training data provides stronger health-related knowledge.")

    add_formatted_paragraph(doc, "C. Evaluation Metric Analysis", font_size=10, bold=True, italic=True,
                            alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=8, space_after=4)

    add_formatted_paragraph(doc,
        "An important methodological finding concerns the divergence between ROUGE-L and factual accuracy metrics. Chain-of-thought and RAG + Optimized prompting produced verbose, structured responses that achieved low ROUGE-L scores despite containing correct factual content. This occurs because ROUGE-L measures longest common subsequence overlap, which penalizes responses that embed correct facts within longer reasoning chains. Our composite factual accuracy metric, weighting keyword recall at 0.6 and ROUGE-L at 0.4, better captures the correctness of verbose but accurate responses. This finding has important implications for hallucination evaluation methodology: relying solely on lexical overlap metrics may systematically underestimate the effectiveness of structured prompting techniques.")

    # ---- V. DISCUSSION ----
    add_formatted_paragraph(doc, "V. Discussion", font_size=10, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=4)

    add_formatted_paragraph(doc,
        "Our results reveal several key insights about hallucination reduction in smaller LLMs. First, the progressive improvement from prompt engineering to RAG confirms the hypothesis that grounding model outputs in external knowledge is more effective than relying on the model's parametric knowledge alone. The 43.3% improvement in factual accuracy from baseline to RAG + Optimized Prompting demonstrates that even smaller models (3B parameters) can achieve meaningful hallucination reduction when augmented with retrieval.")

    add_formatted_paragraph(doc,
        "Second, the relative underperformance of chain-of-thought prompting on smaller models contrasts with findings by Dang et al. [1] on larger models like GPT-4. This suggests that CoT's effectiveness may be capacity-dependent—smaller models may lack the reasoning depth to fully leverage step-by-step decomposition, instead generating plausible-sounding but unfaithful reasoning chains. This aligns with the model variability hypothesis in [1].")

    add_formatted_paragraph(doc,
        "Third, self-consistency decoding showed moderate improvement with minimal implementation complexity, making it a practical first-line defense against hallucinations. The majority-vote mechanism naturally filters out low-confidence responses, though at the cost of 3x inference time.")

    add_formatted_paragraph(doc,
        "The category-level analysis reveals that RAG is particularly effective for factual recall questions (Geography, Health) where retrieved documents directly contain relevant information, but less effective for questions requiring nuanced reasoning about historical or conceptual misconceptions. This suggests that optimal hallucination reduction strategies should be domain-adaptive.")

    add_formatted_paragraph(doc, "A. Limitations", font_size=10, bold=True, italic=True,
                            alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=8, space_after=4)

    add_formatted_paragraph(doc,
        "This study has several limitations. The dataset of 30 questions, while curated to cover diverse categories, is relatively small. The use of a 3B parameter model may not generalize to larger LLMs where techniques like CoT may be more effective. The RAG knowledge base was purpose-built to contain answers to the test questions, representing an optimistic scenario; real-world RAG deployments face additional challenges of retrieval quality and knowledge base completeness. Additionally, our evaluation relies on automated metrics rather than human judgment for all 30 questions, which may miss subtle hallucination patterns.")

    # ---- VI. CONCLUSION ----
    add_formatted_paragraph(doc, "VI. Conclusion and Future Work", font_size=10, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=4)

    add_formatted_paragraph(doc,
        f"This paper presented a systematic evaluation of hallucination reduction techniques for LLaMA 3.2 (3B parameters). Our three-phase experimental pipeline—baseline measurement, prompt-based reduction, and RAG-based reduction—demonstrated that Retrieval-Augmented Generation combined with optimized prompting achieves the highest factual accuracy improvement of {improvement:.1f}% over the vanilla baseline. Among prompt-only methods, few-shot prompting and self-consistency decoding showed the most consistent improvements, while chain-of-thought prompting was less effective on this smaller model.")

    add_formatted_paragraph(doc,
        "Future work will extend this evaluation to larger models (7B, 13B, 70B parameters) to investigate scale-dependent effectiveness. We plan to implement hallucination-aware fine-tuning following the RAG-HAT approach [4] and evaluate on larger benchmarks including the full TruthfulQA dataset. Additionally, we will explore hybrid approaches combining self-consistency with RAG and investigate the use of Natural Language Inference models for automated faithfulness evaluation.")

    # ---- REFERENCES ----
    add_formatted_paragraph(doc, "References", font_size=10, bold=True,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=4)

    references = [
        '[1] A.-H. Dang, T. Vu, and L.-M. Nguyen, "A survey and analysis of hallucinations in large language models, probing whether prompting strategies or model behavior drives them," Frontiers in Artificial Intelligence, vol. 8, 2025.',
        '[2] Y. Bang et al., "HalluLens: A benchmark for LLM hallucinations," in Proc. 63rd ACL, 2025.',
        '[3] X. Wang et al., "Self-consistency improves chain of thought reasoning in language models," in Proc. ICLR, 2023.',
        '[4] J. Song et al., "RAG-HAT: A Hallucination-Aware Tuning Pipeline for LLM in Retrieval-Augmented Generation," in Proc. EMNLP, 2024.',
        '[5] S. Lin, J. Hilton, and O. Evans, "TruthfulQA: Measuring how models mimic human falsehoods," in Proc. ACL, 2022.',
        '[6] C.-Y. Lin, "ROUGE: A package for automatic evaluation of summaries," in Text Summarization Branches Out, 2004.',
        '[7] P. Lewis et al., "Retrieval-augmented generation for knowledge-intensive NLP tasks," in Proc. NeurIPS, 2020.',
        '[8] J. Wei et al., "Chain-of-thought prompting elicits reasoning in large language models," in Proc. NeurIPS, 2022.',
    ]

    for ref in references:
        add_formatted_paragraph(doc, ref, font_size=9, space_after=3)

    # Set the last section (body content) to 2 columns
    last_section = doc.sections[-1]
    last_section.page_width = Inches(8.5)
    last_section.page_height = Inches(11)
    last_section.top_margin = Inches(0.75)
    last_section.bottom_margin = Inches(1)
    last_section.left_margin = Inches(0.625)
    last_section.right_margin = Inches(0.625)
    set_section_columns(last_section, num_cols=2, spacing_twips=360)

    # Save
    output_path = os.path.join(OUTPUT_DIR, "AI700-001-IEEE-Hallucination-LLM-Paper.docx")
    doc.save(output_path)
    print(f"IEEE paper saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    generate_paper()

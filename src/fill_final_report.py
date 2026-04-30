"""Fill the IEEE Word template with the final report content.

Strategy: open the (LibreOffice/Word resaved) clean template, modify paragraphs
in-place by clearing runs and rewriting their text + style. Preserves the
template's multi-column section breaks. Inserts figures inline.
"""
import json
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parent.parent
TEMPLATE = ROOT / "AI700-IEEE-template-clean.docx"
OUT = ROOT / "AI700-Final-Report-Hallucination-LLM.docx"

STATS = json.loads((ROOT / "results" / "stats_summary.json").read_text())
ERR = json.loads((ROOT / "results" / "error_analysis.json").read_text())


def set_para_text(p, text, style=None):
    """Clear runs/hyperlinks and set text. Preserve paragraph properties (pPr)."""
    # Remove anything that could hold visible text: w:r, w:hyperlink, w:smartTag, w:fldSimple
    text_holding_tags = {qn("w:r"), qn("w:hyperlink"), qn("w:smartTag"),
                         qn("w:fldSimple"), qn("w:bookmarkStart"), qn("w:bookmarkEnd")}
    for child in list(p._element):
        if child.tag in text_holding_tags:
            p._element.remove(child)
    if style:
        p.style = style
    if text:
        p.add_run(text)
    return p


def insert_para_after(reference_p, text, style):
    """Insert a new paragraph after the reference paragraph with given style."""
    new_p = deepcopy(reference_p._element)
    # clear children but keep pPr
    for child in list(new_p):
        if child.tag != qn("w:pPr"):
            new_p.remove(child)
    reference_p._element.addnext(new_p)
    from docx.text.paragraph import Paragraph
    np = Paragraph(new_p, reference_p._parent)
    np.style = style
    if text:
        np.add_run(text)
    return np


# ---------- Content ----------

TITLE = "Hallucination Reduction in Large Language Models: A Comparative Study of Prompt Engineering and Retrieval-Augmented Generation"

AUTHOR_LINES = [
    "Emilio Garcia Martinez",
    "Department of Artificial Intelligence",
    "AI700-001 | Student ID: 100783029",
    "Professor: Reda Nacif Elalaoui",
]

ABSTRACT = (
    "One of the biggest challenges with today's Large Language Models is their tendency "
    "to hallucinate—producing text that sounds confident and well-written but is "
    "factually wrong. In this work, we set out to explore how different techniques can "
    "help reduce this problem. We tested LLaMA 3.2 (3B parameters) on a set of 30 "
    "questions drawn from TruthfulQA, covering topics like health, science, history, "
    "geography, and common misconceptions. We tried six different approaches: a plain "
    "vanilla baseline, chain-of-thought prompting, few-shot prompting, self-consistency "
    "decoding, a Retrieval-Augmented Generation (RAG) pipeline, and a combined RAG "
    "approach with optimized prompts. What we found was encouraging—the RAG-based "
    "methods improved factual accuracy by up to 43.3% compared to the baseline, going "
    "from a score of 0.388 to 0.556. For the final phase of this work we add bootstrap "
    "95% confidence intervals (10,000 resamples) and paired bootstrap tests against the "
    "baseline; RAG + Optimized produces a +0.168 paired improvement at p=0.001, with RAG "
    "and self-consistency also reaching statistical significance. A qualitative error "
    "analysis shows that the residual failures across all methods are dominated by "
    "partial-truth responses rather than wholesale fabrication. Overall, giving the "
    "model access to a curated knowledge base made a much bigger difference than just "
    "tweaking how we asked the questions."
)

KEYWORDS = (
    "Large Language Models, Hallucination Reduction, Retrieval-Augmented Generation, "
    "Prompt Engineering, Chain-of-Thought, LLaMA, Factual Accuracy"
)

# Section content (heading_text, style, [body paragraphs and their styles])
# Body items are tuples: (text, style)
SECTIONS = [
    # ---------------- I. INTRODUCTION ----------------
    ("Introduction", "Heading 1", [
        ("Over the past couple of years, Large Language Models like GPT-4, LLaMA, and "
         "Claude have gotten remarkably good at understanding and generating human "
         "language. They are being used in all sorts of areas—healthcare, law, "
         "education, coding—and the results can be genuinely impressive [1]. But there "
         "is a catch. These models have a well-documented habit of making things up. "
         "They produce responses that read perfectly well, but the facts in them can be "
         "completely fabricated [2]. When you are dealing with something like medical "
         "information or legal references, that kind of mistake is not just annoying—it "
         "can be dangerous.",
         "Body Text"),
        ("Researchers have been working on this problem from a few different angles. "
         "Some have focused on smarter prompting strategies, like asking the model to "
         "think step-by-step or giving it examples of correct answers to follow [1]. "
         "Others have tried generating multiple answers and picking the one that comes "
         "up most often [3]. And then there is the RAG approach, which basically lets "
         "the model look things up in an external knowledge base before answering [4]. "
         "Each of these ideas has shown promise on its own, but not many studies have "
         "actually lined them all up and compared them head-to-head on the same model "
         "with the same questions.",
         "Body Text"),
        ("That is exactly what we set out to do here. We took LLaMA 3.2, a smaller "
         "3B-parameter model, and ran it through six different experimental conditions "
         "on a carefully chosen set of 30 questions. The questions span five different "
         "categories—Health, Science, History, Geography, and Misconceptions—so we "
         "could see whether certain techniques work better for certain types of "
         "knowledge. For the final phase of this work we go a step further: we add "
         "bootstrap-based confidence intervals and paired significance tests against "
         "the baseline, and we categorize the model's wrong answers by failure mode so "
         "we can see what kinds of mistakes survive each mitigation.",
         "Body Text"),
    ]),

    # ---------------- II. RELATED WORK ----------------
    ("Related Work", "Heading 1", [
        ("A recent study by Dang, Vu, and Nguyen [1] caught our attention because they "
         "tried to figure out whether hallucinations come from bad prompts or from the "
         "model itself. They came up with two metrics—Prompt Sensitivity and Model "
         "Variability—and tested them on GPT-4, LLaMA 2, and DeepSeek. Their takeaway "
         "was interesting: chain-of-thought prompting helps a lot when the "
         "hallucination is prompt-related, but some models just hallucinate no matter "
         "what you do. That told us we would probably need more than just better "
         "prompts.",
         "Body Text"),
        ("On the evaluation side, Bang et al. [2] built HalluLens, which gives you a "
         "structured way to test for different kinds of hallucinations—things the model "
         "makes up versus things it contradicts itself on. We drew on their "
         "categorization when designing our own evaluation approach. Song et al. [4] "
         "took a different route with RAG-HAT, where they actually fine-tuned a model "
         "to be aware of its own hallucinations using a three-stage pipeline. Their "
         "results were pretty compelling and motivated our decision to include RAG in "
         "our experiments. Wang et al. [3] showed that if you just ask the model the "
         "same question a few times and go with the most common answer, you tend to "
         "get better results—a simple but surprisingly effective trick.",
         "Body Text"),
    ]),

    # ---------------- III. METHODOLOGY ----------------
    ("Methodology", "Heading 1", [
        ("Experimental Setup", "Heading 2"),
        ("We ran all our experiments on LLaMA 3.2 with 3 billion parameters, hosted "
         "locally through Ollama on an Apple M1 machine with 8 GB of RAM. The "
         "temperature was set to 0.7 and we capped responses at 300 tokens. For our "
         "test set, we put together 30 questions inspired by the TruthfulQA benchmark "
         "[5]. Each question had a verified correct answer and a couple of known wrong "
         "answers that models commonly give. We spread the questions across five "
         "categories: Health (7 questions), Science (8), History (5), Geography (4), "
         "and Misconceptions (6).",
         "Body Text"),
        ("Evaluation Metrics", "Heading 2"),
        ("We used three different ways to evaluate the model's answers. First, a "
         "hallucination detection classifier that checks whether the response lines up "
         "more with the correct answer or the known wrong ones, using a mix of "
         "keyword matching and ROUGE-L similarity. Second, the ROUGE-L score itself "
         "[6], which measures how much overlap there is between what the model said "
         "and the ground truth. Third, a factual accuracy score we designed "
         "ourselves—it combines keyword recall (weighted at 0.6) with ROUGE-L "
         "(weighted at 0.4). We needed this custom metric because methods like "
         "chain-of-thought produce long, wordy answers that get penalized by ROUGE-L "
         "even when the actual facts are right.",
         "Body Text"),
        ("Phase 1: Baseline Measurement", "Heading 2"),
        ("For the baseline, we just asked each question with a straightforward prompt: "
         "\"Answer the following question concisely and factually.\" No tricks, no "
         "special formatting. This gave us a starting point to measure everything "
         "else against.",
         "Body Text"),
        ("Phase 2: Prompt-Based Reduction", "Heading 2"),
        ("We tested three different prompting strategies. With Chain-of-Thought, we "
         "told the model to reason through the question step by step, think about what "
         "misconceptions exist, and check against scientific evidence before giving a "
         "final answer. For Few-Shot, we included three worked examples of accurate, "
         "myth-busting answers right in the prompt. Self-Consistency was a bit "
         "different—we ran each question three times with a slightly higher "
         "temperature (0.8) and then picked whichever answer was most similar to the "
         "others, essentially letting the model vote with itself.",
         "Body Text"),
        ("Phase 3: RAG-Based Reduction", "Heading 2"),
        ("We built a small knowledge base from four reference documents covering "
         "health myths, scientific misconceptions, historical facts, and geography. "
         "These were chunked into paragraphs and indexed using ChromaDB with the "
         "all-MiniLM-L6-v2 sentence embedding model, giving us 27 searchable chunks "
         "in total. We tested two variants: a basic RAG setup that retrieves the top 3 "
         "relevant chunks and tells the model to answer based on them, and an "
         "optimized version that adds chain-of-thought instructions, a few-shot "
         "example, and explicit fact-checking guidance on top of the retrieved "
         "context.",
         "Body Text"),
        ("Statistical Analysis", "Heading 2"),
        ("For the final phase we add a statistical layer on top of the existing "
         "results. Per-method means are reported with 95% confidence intervals "
         "computed by non-parametric bootstrap with 10,000 resamples. Paired "
         "comparisons against the baseline use the same questions for each method; "
         "the bootstrap is performed on the per-question difference, and a two-sided "
         "p-value is computed as twice the fraction of bootstrap means with sign "
         "opposite to the point estimate.",
         "Body Text"),
    ]),

    # ---------------- IV. RESULTS ----------------
    ("Results", "Heading 1", [
        ("Overall Performance", "Heading 2"),
        ("__TABLE_PERF__", None),
        ("The results in Table I tell a clear story. Factual accuracy climbed steadily "
         "as we moved from simple prompting to retrieval-based methods. The baseline "
         "scored 0.388, and by the time we got to RAG with optimized prompts, that had "
         "jumped to 0.556—a 43.3% improvement. The plain RAG pipeline actually had the "
         "best ROUGE-L score at 0.353, which makes sense since it pulls in relevant "
         "text that the model can echo back.",
         "Body Text"),
        ("What surprised us a bit was how the prompt-only methods performed. Few-shot "
         "did reasonably well, pushing accuracy up to 0.460. Self-consistency managed "
         "to bring the hallucination rate down slightly, from 66.7% to 70.0%. But "
         "chain-of-thought was a disappointment—its ROUGE-L score tanked to 0.102 "
         "because the model rambled through long reasoning chains. The facts were "
         "sometimes in there, but buried under so much extra text that our metrics "
         "struggled to find them.",
         "Body Text"),
        ("Statistical Significance", "Heading 2"),
        ("Adding the bootstrap analysis lets us be more precise about which "
         "improvements are real and which could be noise on a 30-question set. Three "
         "methods clear the 5% significance bar in a paired test against the baseline: "
         "RAG + Optimized (Δ=+0.168, 95% CI [0.064, 0.270], p=0.001), RAG (Δ=+0.121, "
         "[0.010, 0.228], p=0.031) and Self-Consistency (Δ=+0.060, [0.005, 0.119], "
         "p=0.034). Few-Shot shows a positive trend that does not reach significance "
         "(Δ=+0.072, p=0.125), and Chain-of-Thought is indistinguishable from baseline "
         "(Δ=–0.042, p=0.340). Fig. 1 shows the per-method accuracy with confidence "
         "intervals; Fig. 2 shows the paired differences with significance markers.",
         "Body Text"),
        ("__FIGURE__:results/accuracy_with_ci.png:Factual accuracy by method "
         "with bootstrap 95% confidence intervals.",
         "figure caption"),
        ("__FIGURE__:results/paired_diffs.png:Paired accuracy improvement "
         "over the baseline. * p<0.05, ** p<0.01, *** p<0.001.",
         "figure caption"),
        ("Category-Level Analysis", "Heading 2"),
        ("When we broke things down by category, the picture got more nuanced. "
         "Geography questions saw the biggest boost from RAG, which makes intuitive "
         "sense—facts about capitals and deserts are exactly the kind of thing a "
         "knowledge base excels at. Health questions already had a relatively low "
         "hallucination rate at baseline, probably because the model picked up decent "
         "medical knowledge during pre-training. The tough categories were History "
         "and Misconceptions. These often require the model to not just recall a "
         "fact, but to recognize and correct a common misunderstanding. That turned "
         "out to be harder than simply looking something up.",
         "Body Text"),
        ("Evaluation Metric Analysis", "Heading 2"),
        ("One thing we did not expect was how much ROUGE-L and our factual accuracy "
         "metric would disagree. Chain-of-thought and RAG + Optimized both generated "
         "longer, more structured answers. ROUGE-L punished them for it, since there "
         "is less direct word overlap with a short ground-truth answer when your "
         "response is three paragraphs long. But when we looked at keyword "
         "recall—whether the key facts actually showed up in the response—the picture "
         "was much better. This is why we weighted keyword recall at 0.6 in our "
         "composite score. It is a reminder that picking the right evaluation metric "
         "really matters, especially when you are comparing methods that produce very "
         "different output styles.",
         "Body Text"),
        ("Error-Type Analysis", "Heading 2"),
        ("To dig deeper into the residual failures, for the final phase we "
         "categorized every non-correct response into one of four buckets using "
         "lightweight heuristics: partial-truth (substantial overlap with the correct "
         "answer but still flagged as hallucination), fabrication (no overlap, "
         "declarative), overconfident-fabrication (long and declarative with no "
         "hedging), and refusal-or-hedge. Fig. 3 shows the breakdown by method.",
         "Body Text"),
        ("__FIGURE__:results/error_types.png:Response-type distribution by "
         "method. Most non-correct responses are partial-truths rather than full "
         "fabrications.",
         "figure caption"),
        ("Two patterns jump out. First, the dominant failure mode across every method "
         "is partial-truth: the model surfaces the right topic and even some of the "
         "correct vocabulary, but does not commit to the verified answer. That tells "
         "us the next ceiling is not going to come from surfacing more facts—RAG "
         "already does that—but from helping the model pick the right one and stick "
         "with it. Second, even with RAG and RAG + Optimized we still see a small "
         "number of overconfident fabrications where the model invents details around "
         "the retrieved context, which lines up with the kind of failure that "
         "hallucination-aware fine-tuning [4] is designed to address.",
         "Body Text"),
    ]),

    # ---------------- V. DISCUSSION ----------------
    ("Discussion", "Heading 1", [
        ("Looking at the big picture, the clearest lesson from our experiments is "
         "that giving the model something to reference makes a real difference. "
         "Prompt engineering on its own can help, but it is fighting with one hand "
         "tied behind its back—the model can only work with what it already \"knows\" "
         "from training, and for a 3B-parameter model, that knowledge has gaps. RAG "
         "fills those gaps by putting verified information right in front of the "
         "model. Importantly, the bootstrap analysis we added in the final phase "
         "confirms this is not just noise: the +0.168 paired gain for RAG + "
         "Optimized has a 95% CI that does not cross zero, and the two-sided "
         "p-value is 0.001.",
         "Body Text"),
        ("The chain-of-thought results were honestly a bit frustrating. On paper, "
         "asking a model to reason step-by-step should help it catch its own "
         "mistakes. And for bigger models like GPT-4, Dang et al. [1] showed that it "
         "does. But our 3B model seemed to use the extra reasoning steps to "
         "construct more elaborate wrong answers rather than to correct itself. The "
         "error-type analysis backs this up—Chain-of-Thought has the highest share "
         "of refusal-or-hedge responses, suggesting the model exposes its own "
         "uncertainty without converging on an answer. It is a good reminder that "
         "techniques developed on frontier models do not always scale down "
         "gracefully.",
         "Body Text"),
        ("Self-consistency turned out to be the most practical prompt-only "
         "technique, and the only one besides RAG to clear statistical "
         "significance. It does not require any special prompt design or external "
         "resources—you just run the query a few times and pick the consensus. The "
         "downside is that it triples your inference cost, which matters if you are "
         "running on consumer hardware like we were.",
         "Body Text"),
        ("One pattern that stood out in the category analysis is that RAG works "
         "best when the answer is a straightforward fact that can be looked up, "
         "and worse when the question requires the model to reason about why a "
         "common belief is wrong. That suggests you might want different strategies "
         "for different types of questions in a production system.",
         "Body Text"),
        ("Limitations", "Heading 2"),
        ("We want to be upfront about the limitations of this work. Thirty "
         "questions is a small dataset, and while the bootstrap analysis says the "
         "RAG + Optimized improvement is unlikely to be a chance result, we are "
         "still working with a tight statistical envelope. We only tested one "
         "model size—3B parameters—and our RAG knowledge base was specifically "
         "built to contain the answers, which is more optimistic than what you "
         "would get in the real world where retrieval quality varies. We also "
         "relied entirely on automated metrics. Having human evaluators review the "
         "responses would give us a more complete picture, and that is something "
         "we plan to add in future work.",
         "Body Text"),
    ]),

    # ---------------- VI. CONCLUSION ----------------
    ("Conclusion and Future Work", "Heading 1", [
        ("In this paper, we walked through a three-phase experiment to see what "
         "actually works for reducing hallucinations in a smaller language model. "
         "The bottom line is that retrieval-augmented generation, especially when "
         "paired with well-designed prompts, gave us the best results—a 43.3% "
         "improvement in factual accuracy over the plain baseline, and a paired "
         "bootstrap gain of +0.168 at p=0.001. Prompt engineering by itself helped "
         "somewhat, with few-shot and self-consistency being the most reliable "
         "techniques, but it was not enough on its own to make a dramatic "
         "difference at this model size.",
         "Body Text"),
        ("The error-type analysis we added in the final phase tells us where to go "
         "next. Most of the residual failures are partial-truths rather than full "
         "fabrications, which means the model usually has the right region of "
         "knowledge but does not commit to the verified answer. Going forward, we "
         "want to try these same experiments on larger models to see if techniques "
         "like chain-of-thought start working better with more parameters, and we "
         "are interested in hallucination-aware fine-tuning along the lines of "
         "what Song et al. [4] proposed with RAG-HAT. Running the evaluation on "
         "the full TruthfulQA benchmark rather than our smaller subset would make "
         "the results more generalizable. And finally, we would like to explore "
         "using Natural Language Inference models to automatically assess whether "
         "the model's answer is actually faithful to the retrieved context, "
         "rather than relying on surface-level text matching.",
         "Body Text"),
    ]),
]

ACK_HEADING = "Acknowledgment"
ACK_BODY = (
    "The author thanks Prof. Reda Nacif Elalaoui for the AI700-001 course and the "
    "structure that motivated this work, and the open-source maintainers of Ollama, "
    "ChromaDB and the TruthfulQA dataset, without which a local, reproducible "
    "evaluation would not have been possible."
)

REFERENCES = [
    "A.-H. Dang, T. Vu, and L.-M. Nguyen, \"A survey and analysis of hallucinations in large language models,\" Frontiers in Artificial Intelligence, vol. 8, 2025.",
    "Y. Bang et al., \"HalluLens: A benchmark for LLM hallucinations,\" in Proc. 63rd ACL, 2025.",
    "X. Wang et al., \"Self-consistency improves chain of thought reasoning in language models,\" in Proc. ICLR, 2023.",
    "J. Song et al., \"RAG-HAT: A Hallucination-Aware Tuning Pipeline for LLM in Retrieval-Augmented Generation,\" in Proc. EMNLP, 2024.",
    "S. Lin, J. Hilton, and O. Evans, \"TruthfulQA: Measuring how models mimic human falsehoods,\" in Proc. ACL, 2022.",
    "C.-Y. Lin, \"ROUGE: A package for automatic evaluation of summaries,\" in Text Summarization Branches Out, 2004.",
    "P. Lewis et al., \"Retrieval-augmented generation for knowledge-intensive NLP tasks,\" in Proc. NeurIPS, 2020.",
    "J. Wei et al., \"Chain-of-thought prompting elicits reasoning in large language models,\" in Proc. NeurIPS, 2022.",
]


# ---------- Document construction ----------

def main():
    """Build the report fresh: keep the IEEE template's *styles*, but clear its
    body and rebuild a clean structure. This avoids the template's nested
    multi-column section breaks that were reordering content."""
    doc = Document(str(TEMPLATE))

    # 0. Find the final page-level sectPr (the one at the end of the body).
    # We keep that to preserve page size and margins; everything else goes.
    body_xml = doc.element.body
    final_sectPr = None
    for child in list(body_xml):
        if child.tag == qn("w:sectPr"):
            final_sectPr = deepcopy(child)
    if final_sectPr is None:
        # Fall back to scraping a sectPr embedded in the last paragraph
        for p in reversed(list(body_xml)):
            sp = p.find(qn("w:pPr/w:sectPr")) if p.tag == qn("w:p") else None
            if sp is not None:
                final_sectPr = deepcopy(sp)
                break

    # Make the final section a clean 2-column body (IEEE style)
    if final_sectPr is not None:
        cols = final_sectPr.find(qn("w:cols"))
        if cols is None:
            from lxml import etree
            cols = etree.SubElement(final_sectPr, qn("w:cols"))
        cols.set(qn("w:num"), "2")
        cols.set(qn("w:space"), "36pt")
        # Remove column-only "type continuous" so this acts as the page-level section
        type_el = final_sectPr.find(qn("w:type"))
        if type_el is not None:
            final_sectPr.remove(type_el)

    # 1. Clear all body children (paragraphs, tables, section breaks)
    for child in list(body_xml):
        body_xml.remove(child)

    # Helper to add a paragraph with style
    def add_para(text, style):
        p = doc.add_paragraph()
        if style:
            try:
                p.style = style
            except KeyError:
                pass
        if text:
            p.add_run(text)
        return p

    # 2. Title section (single column)
    # Insert a section break BEFORE the title-area sectPr so title is not in 2-col flow
    # Strategy: title + author + abstract all live in a single-column section,
    # then a section break, then 2-column body.
    title_p = add_para(TITLE, "paper title")

    # Author block - one paragraph, centered (style "Author")
    author_p = add_para("\n".join(AUTHOR_LINES), "Author")
    # Add explicit line breaks so the runs render as separate lines
    # (Author style centers and python-docx already preserves \n if we use add_run with breaks)
    # Replace newlines with <w:br/>
    for run in list(author_p.runs):
        run._element.getparent().remove(run._element)
    for i, line in enumerate(AUTHOR_LINES):
        if i > 0:
            r = author_p.add_run()
            br = r._element.makeelement(qn("w:br"), {})
            r._element.append(br)
        author_p.add_run(line)

    # 3. Abstract + Keywords go in the SAME single-column section as title+author,
    # so they sit directly under the author block (no page break) — matches the
    # midterm paper's layout.
    abstract_p = add_para("Abstract—" + ABSTRACT, "Abstract")
    keywords_p = add_para("Keywords—" + KEYWORDS, "Keywords")

    # Now insert a continuous section break that switches to 2-column for the
    # body content that follows.
    from lxml import etree
    sect_break_p = doc.add_paragraph()
    pPr = sect_break_p._element.find(qn("w:pPr"))
    if pPr is None:
        pPr = etree.SubElement(sect_break_p._element, qn("w:pPr"))
        sect_break_p._element.insert(0, pPr)
    sectPr_inline = etree.SubElement(pPr, qn("w:sectPr"))
    type_el = etree.SubElement(sectPr_inline, qn("w:type"))
    type_el.set(qn("w:val"), "continuous")
    cols_el = etree.SubElement(sectPr_inline, qn("w:cols"))
    cols_el.set(qn("w:num"), "1")
    cols_el.set(qn("w:space"), "36pt")

    # Mock vars to keep downstream code working
    paras = doc.paragraphs
    orig_paras = list(paras)
    # We deliberately skip the old in-place fill path; jump straight to body content.

    # ----- Add main body content with proper section numbering -----
    figure_inserts = []
    table_inserts = []

    # The IEEE template's Heading 1 and Heading 2 styles already auto-number
    # (Roman numerals for H1, capital letters for H2). So we just pass the
    # plain heading text and let the style do the numbering.
    for heading, h_style, items in SECTIONS:
        add_para(heading, h_style)
        for txt, st in items:
            if isinstance(txt, str) and txt.startswith("__FIGURE__:"):
                _, img_path, caption = txt.split(":", 2)
                # We render: image paragraph, then caption paragraph
                img_p = doc.add_paragraph()
                img_p.alignment = 1  # center
                run = img_p.add_run()
                full_path = ROOT / img_path
                if full_path.exists():
                    try:
                        run.add_picture(str(full_path), width=Inches(3.3))
                    except Exception as e:
                        print(f"  ! failed to add image {full_path}: {e}")
                add_para(caption, "figure caption")
            elif txt == "__TABLE_PERF__":
                # We'll add a table here directly
                add_para("Per-method performance with bootstrap 95% confidence intervals (10,000 resamples).", "table head")
                table = doc.add_table(rows=1, cols=4)
                hdr = table.rows[0].cells
                hdr[0].text = "Method"
                hdr[1].text = "Factual Acc. (95% CI)"
                hdr[2].text = "Hallucination Rate"
                hdr[3].text = "ROUGE-L"
                method_labels = {
                    "baseline": "Baseline",
                    "cot": "Chain-of-Thought",
                    "few_shot": "Few-Shot",
                    "self_consistency": "Self-Consistency",
                    "rag": "RAG",
                    "rag_optimized": "RAG + Optimized",
                }
                for m in ["baseline", "cot", "few_shot", "self_consistency", "rag", "rag_optimized"]:
                    row = table.add_row().cells
                    a = STATS[m]["accuracy"]
                    h = STATS[m]["hallucination_rate"]
                    r = STATS[m]["rouge_l"]
                    row[0].text = method_labels[m]
                    row[1].text = f"{a['mean']:.3f} [{a['ci_lo']:.3f}, {a['ci_hi']:.3f}]"
                    row[2].text = f"{h['mean']:.3f}"
                    row[3].text = f"{r['mean']:.3f}"
            elif st == "Heading 2":
                add_para(txt, "Heading 2")
            else:
                add_para(txt, st or "Body Text")

    # ----- Acknowledgment -----
    add_para(ACK_HEADING, "Heading 5")
    add_para(ACK_BODY, "Body Text")

    # ----- References -----
    add_para("References", "Heading 5")
    for ref in REFERENCES:
        add_para(ref, "references")

    # 5. Re-attach the page-level sectPr at the very end of body
    if final_sectPr is not None:
        body_xml.append(final_sectPr)

    doc.save(str(OUT))
    print(f"Wrote {OUT}")
    return  # short-circuit; legacy fill path below is unused

    # 2. Authors: collapse the multi-author block into one filled author + blanks.
    # The Author paragraphs are 1..18.  We place one filled author at paragraph 3
    # (which corresponds to the first 4-column author cell), and clear the rest.
    author_text = "\n".join(AUTHOR_LINES)
    for i in range(1, 19):
        if i == 3:
            set_para_text(paras[i], author_text, "Author")
        else:
            set_para_text(paras[i], "", "Author")

    # 3. Abstract / Keywords
    set_para_text(paras[19], "Abstract—" + ABSTRACT, "Abstract")
    set_para_text(paras[20], "Keywords—" + KEYWORDS, "Keywords")

    # 4. Body. Paragraphs 21..79 are body-content slots before the original
    # References Heading-5 at index 80. We rewrite them in place and remove
    # any unused slots.
    BODY_START = 21
    BODY_END = 79
    body_slots = list(range(BODY_START, BODY_END + 1))

    # Flatten our content into (text, style) tuples
    flat = []
    for heading, h_style, items in SECTIONS:
        flat.append((heading, h_style))
        for txt, st in items:
            flat.append((txt, st))
    flat.append((ACK_HEADING, "Heading 5"))
    flat.append((ACK_BODY, "Body Text"))

    # Process __FIGURE__ and __TABLE_PERF__ markers separately:
    # Render figures by clearing the slot paragraph, then inserting an image
    # before the caption text.

    # We also need to track where each slot ends up so we can insert the table.
    body_iter = iter(body_slots)
    figure_inserts = []  # list of (paragraph_obj, image_path)
    table_inserts = []   # list of (paragraph_obj, kind)

    # First pass: assign content to existing slots
    for entry in flat:
        try:
            slot = next(body_iter)
        except StopIteration:
            # Need to add a new paragraph at end of body section
            # Append after BODY_END paragraph by creating new paragraphs
            ref = paras[BODY_END]
            text, style = entry
            if isinstance(text, str) and text.startswith("__FIGURE__:"):
                _, img_path, caption = text.split(":", 2)
                np = insert_para_after(paras[BODY_END], "", "figure caption")
                figure_inserts.append((np, img_path, caption))
            elif text == "__TABLE_PERF__":
                np = insert_para_after(paras[BODY_END], "", "Body Text")
                table_inserts.append((np, "perf"))
            else:
                np = insert_para_after(paras[BODY_END], text, style or "Body Text")
            paras = doc.paragraphs  # refresh
            continue

        text, style = entry
        p = paras[slot]
        if isinstance(text, str) and text.startswith("__FIGURE__:"):
            _, img_path, caption = text.split(":", 2)
            set_para_text(p, "", "figure caption")
            figure_inserts.append((p, img_path, caption))
        elif text == "__TABLE_PERF__":
            set_para_text(p, "", "Body Text")
            table_inserts.append((p, "perf"))
        else:
            set_para_text(p, text, style or "Body Text")

    # Remove any remaining body slots entirely so the document doesn't have a
    # large run of empty paragraphs between the Acknowledgment and References.
    for slot in body_iter:
        p = paras[slot]
        p._element.getparent().remove(p._element)

    # 5. Replace references. Use the *original* paras list (captured before any
    # XML removals) — element references remain valid because XML elements
    # aren't moved, only siblings are removed.
    # In the original template:
    #   index 80     = Heading 5 "References"        (keep)
    #   index 81..85 = Body Text guidance about refs (clear)
    #   index 86..97 = sample reference entries      (replace with our refs)
    if 80 < len(orig_paras):
        set_para_text(orig_paras[80], "References", "Heading 5")
    for i in range(81, 86):
        if i < len(orig_paras):
            p = orig_paras[i]
            p._element.getparent().remove(p._element)

    REF_START = 86
    REF_END = 97
    ref_slots = list(range(REF_START, REF_END + 1))
    for i, ref in enumerate(REFERENCES):
        if i < len(ref_slots) and ref_slots[i] < len(orig_paras):
            # The "references" style auto-numbers, so don't add a [n] prefix.
            set_para_text(orig_paras[ref_slots[i]], ref, "references")
    # Remove any unused reference slots
    for j in range(len(REFERENCES), len(ref_slots)):
        if ref_slots[j] < len(orig_paras):
            p = orig_paras[ref_slots[j]]
            p._element.getparent().remove(p._element)

    # 6. Insert figures: replace caption paragraph with image-then-caption
    for fig_p, img_path, caption in figure_inserts:
        full_path = ROOT / img_path
        if not full_path.exists():
            print(f"  ! missing image: {full_path}")
            set_para_text(fig_p, caption, "figure caption")
            continue
        # Insert a new paragraph BEFORE the caption holding the image
        from docx.text.paragraph import Paragraph
        new_p_xml = deepcopy(fig_p._element)
        # clear runs in new paragraph
        for child in list(new_p_xml):
            if child.tag == qn("w:r"):
                new_p_xml.remove(child)
        fig_p._element.addprevious(new_p_xml)
        img_p = Paragraph(new_p_xml, fig_p._parent)
        img_p.style = "figure caption"
        run = img_p.add_run()
        try:
            run.add_picture(str(full_path), width=Inches(3.3))
        except Exception as e:
            print(f"  ! failed to add image {full_path}: {e}")
        # Set caption text
        set_para_text(fig_p, caption, "figure caption")

    # 7. Insert performance table
    for tbl_p, kind in table_inserts:
        if kind == "perf":
            # Build a 7-row x 4-col table just before the paragraph
            table = doc.add_table(rows=1, cols=4)
            try:
                table.style = "Table Grid"
            except KeyError:
                pass  # template lacks Table Grid; use default
            hdr = table.rows[0].cells
            hdr[0].text = "Method"
            hdr[1].text = "Factual Acc. (95% CI)"
            hdr[2].text = "Hallucination Rate"
            hdr[3].text = "ROUGE-L"
            method_labels = {
                "baseline": "Baseline",
                "cot": "Chain-of-Thought",
                "few_shot": "Few-Shot",
                "self_consistency": "Self-Consistency",
                "rag": "RAG",
                "rag_optimized": "RAG + Optimized",
            }
            for m in ["baseline", "cot", "few_shot", "self_consistency", "rag", "rag_optimized"]:
                row = table.add_row().cells
                a = STATS[m]["accuracy"]
                h = STATS[m]["hallucination_rate"]
                r = STATS[m]["rouge_l"]
                row[0].text = method_labels[m]
                row[1].text = f"{a['mean']:.3f} [{a['ci_lo']:.3f}, {a['ci_hi']:.3f}]"
                row[2].text = f"{h['mean']:.3f}"
                row[3].text = f"{r['mean']:.3f}"
            # Move the table to be just before the caption paragraph
            tbl_xml = table._element
            tbl_p._element.addprevious(tbl_xml)
            # Add a Table I caption above the table by repurposing tbl_p's role:
            # we want: caption paragraph (above the table) that says "Table I"
            # For simplicity, re-purpose tbl_p as caption
            set_para_text(tbl_p, "Per-method performance with bootstrap 95% confidence intervals (10,000 resamples).", "table head")
            # Move caption above the table
            tbl_p._element.addprevious(deepcopy(tbl_p._element))
            tbl_p._element.getparent().remove(tbl_p._element)

    doc.save(str(OUT))
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()

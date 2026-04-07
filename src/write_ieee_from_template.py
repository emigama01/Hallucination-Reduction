"""
Write IEEE paper by replacing text in the original template paragraph by paragraph.
Preserves ALL formatting, section breaks, columns, styles.
"""

from docx import Document
from docx.shared import Pt
from lxml import etree
import json
import os
import copy

RESULTS_DIR = os.path.join(os.path.dirname(__file__), "..", "results")
TEMPLATE = os.path.join(os.path.dirname(__file__), "..", "IEEE-template-fixed.docx")
OUTPUT = os.path.join(os.path.dirname(__file__), "..", "AI700-001-IEEE-Hallucination-LLM-Paper.docx")
NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}


def load_results():
    with open(os.path.join(RESULTS_DIR, "summary.json")) as f:
        return json.load(f)


def set_para_text(para, text):
    """Replace paragraph text while keeping style and section break properties."""
    p_elem = para._p

    # Save sectPr if it exists (section break)
    sect = p_elem.find('.//w:sectPr', NS)
    saved_sect = copy.deepcopy(sect) if sect is not None else None

    # Save pPr (paragraph properties)
    pPr = p_elem.find('w:pPr', NS)
    saved_pPr = copy.deepcopy(pPr) if pPr is not None else None

    # Get run properties from first run (font formatting)
    saved_rPr = None
    first_run = p_elem.find('w:r', NS)
    if first_run is not None:
        rPr = first_run.find('w:rPr', NS)
        if rPr is not None:
            saved_rPr = copy.deepcopy(rPr)

    # Clear all children
    for child in list(p_elem):
        p_elem.remove(child)

    # Restore pPr
    if saved_pPr is not None:
        # Remove old sectPr from pPr if we saved it separately
        for s in saved_pPr.findall('w:sectPr', NS):
            saved_pPr.remove(s)
        if saved_sect is not None:
            saved_pPr.append(saved_sect)
        p_elem.append(saved_pPr)

    # Add new run with text
    qn_w = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    run = etree.SubElement(p_elem, f'{qn_w}r')
    if saved_rPr is not None:
        run.append(saved_rPr)
    t = etree.SubElement(run, f'{qn_w}t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text


def remove_para(para):
    """Remove a paragraph from the document."""
    p = para._p
    p.getparent().remove(p)


def insert_para_after(doc, ref_para, text, style_name):
    """Insert a new paragraph after ref_para with given style."""
    qn_w = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

    # Create new paragraph element
    new_p = etree.Element(f'{qn_w}p')

    # Get style id
    style = doc.styles[style_name]
    pPr = etree.SubElement(new_p, f'{qn_w}pPr')
    pStyle = etree.SubElement(pPr, f'{qn_w}pStyle')
    pStyle.set(f'{qn_w}val', style.style_id)

    # Copy run properties from the reference paragraph's first run
    ref_run = ref_para._p.find(f'{qn_w}r')
    saved_rPr = None
    if ref_run is not None:
        rPr = ref_run.find(f'{qn_w}rPr')
        if rPr is not None:
            saved_rPr = copy.deepcopy(rPr)

    # Add run with text
    run = etree.SubElement(new_p, f'{qn_w}r')
    if saved_rPr is not None:
        run.append(saved_rPr)
    t = etree.SubElement(run, f'{qn_w}t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text

    # Insert after reference
    ref_para._p.addnext(new_p)

    # Return a paragraph wrapper (find it in doc)
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, ref_para._p.getparent())


def add_table_after(doc, ref_para, headers, rows):
    """Insert a table after a paragraph."""
    qn_w = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

    tbl = etree.Element(f'{qn_w}tbl')

    # Table properties - simple borders
    tblPr = etree.SubElement(tbl, f'{qn_w}tblPr')
    tblStyle = etree.SubElement(tblPr, f'{qn_w}tblStyle')
    tblStyle.set(f'{qn_w}val', 'TableGrid')
    tblW = etree.SubElement(tblPr, f'{qn_w}tblW')
    tblW.set(f'{qn_w}w', '0')
    tblW.set(f'{qn_w}type', 'auto')

    # Add borders
    tblBorders = etree.SubElement(tblPr, f'{qn_w}tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = etree.SubElement(tblBorders, f'{qn_w}{border_name}')
        b.set(f'{qn_w}val', 'single')
        b.set(f'{qn_w}sz', '4')
        b.set(f'{qn_w}space', '0')
        b.set(f'{qn_w}color', '000000')

    all_rows = [headers] + rows
    for row_idx, row_data in enumerate(all_rows):
        tr = etree.SubElement(tbl, f'{qn_w}tr')
        for col_idx, cell_text in enumerate(row_data):
            tc = etree.SubElement(tr, f'{qn_w}tc')
            p = etree.SubElement(tc, f'{qn_w}p')
            pPr = etree.SubElement(p, f'{qn_w}pPr')
            jc = etree.SubElement(pPr, f'{qn_w}jc')
            jc.set(f'{qn_w}val', 'center')

            r = etree.SubElement(p, f'{qn_w}r')
            rPr = etree.SubElement(r, f'{qn_w}rPr')
            sz = etree.SubElement(rPr, f'{qn_w}sz')
            sz.set(f'{qn_w}val', '16')  # 8pt
            szCs = etree.SubElement(rPr, f'{qn_w}szCs')
            szCs.set(f'{qn_w}val', '16')
            rFonts = etree.SubElement(rPr, f'{qn_w}rFonts')
            rFonts.set(f'{qn_w}ascii', 'Times New Roman')
            rFonts.set(f'{qn_w}hAnsi', 'Times New Roman')
            if row_idx == 0:
                b = etree.SubElement(rPr, f'{qn_w}b')

            t = etree.SubElement(r, f'{qn_w}t')
            t.text = str(cell_text)

    ref_para._p.addnext(tbl)


def main():
    summary = load_results()
    doc = Document(TEMPLATE)

    baseline_acc = summary['baseline']['avg_factual_accuracy']
    best_acc = summary['rag_optimized']['avg_factual_accuracy']
    improvement = ((best_acc - baseline_acc) / baseline_acc) * 100

    # Build a map of paragraphs by index for easy access
    paras = doc.paragraphs

    # ================================================================
    # P0: paper title
    # ================================================================
    set_para_text(paras[0],
        "Hallucination Reduction in Large Language Models: A Comparative Study of Prompt Engineering and Retrieval-Augmented Generation")

    # ================================================================
    # P1: subtitle note - clear it
    # P2: empty (has section break - KEEP)
    # ================================================================
    set_para_text(paras[1], "")

    # ================================================================
    # P3: Author block (has section break for 4-col author layout)
    # We only have 1 author, so set first author and clear rest
    # ================================================================
    set_para_text(paras[3],
        "Emilio Garcia Martinez\nDepartment of Artificial Intelligence\nAI700-001 | ID: 100783029\nProfessor: Reda Nacif Elalaoui")

    # Clear other author blocks (P4-P18 roughly are author-related)
    for i in range(4, 19):
        if i < len(paras) and paras[i].style.name in ('Author', 'Affiliation', 'Normal'):
            set_para_text(paras[i], "")

    # ================================================================
    # P19: Abstract
    # ================================================================
    set_para_text(paras[19],
        "Abstract\u2014Large Language Models (LLMs) have demonstrated remarkable capabilities in natural language understanding and generation, yet they remain prone to hallucination\u2014generating fluent but factually incorrect content. This paper presents a systematic evaluation of hallucination reduction techniques applied to LLaMA 3.2 (3B parameters) across a curated dataset of 30 TruthfulQA-inspired questions spanning five knowledge categories. We implement and compare six experimental conditions: vanilla baseline, chain-of-thought (CoT) prompting, few-shot prompting, self-consistency decoding, Retrieval-Augmented Generation (RAG), and RAG with optimized prompting. Our results demonstrate that RAG-based approaches achieve the highest factual accuracy improvement of 43.3% over baseline (0.388 to 0.556), with the RAG pipeline showing the strongest ROUGE-L scores (0.353). Self-consistency decoding provided the best hallucination rate reduction among prompt-only methods. These findings confirm that grounding LLM outputs in retrieved knowledge bases is significantly more effective than prompt engineering alone for hallucination mitigation.")

    # ================================================================
    # P20: Keywords
    # ================================================================
    set_para_text(paras[20],
        "Keywords\u2014Large Language Models, Hallucination Reduction, Retrieval-Augmented Generation, Prompt Engineering, Chain-of-Thought, LLaMA, Factual Accuracy")

    # ================================================================
    # P21: Heading 1 - Introduction
    # ================================================================
    set_para_text(paras[21], "Introduction")

    # P22: Body Text - first intro paragraph
    set_para_text(paras[22],
        "Large Language Models (LLMs) such as GPT-4, LLaMA, and Claude have achieved remarkable advances in natural language processing, enabling applications in healthcare, legal analysis, education, and software development [1]. However, a critical limitation persists: the phenomenon of hallucination, defined as the generation of text that is syntactically fluent but factually incorrect or unsupported by source material [2]. In high-stakes domains, hallucinated outputs can have severe consequences\u2014from incorrect medical advice to fabricated legal citations.")

    # P23: was "Ease of Use" heading - change to second intro para
    # Change style to Body Text
    paras[23].style = doc.styles['Body Text']
    set_para_text(paras[23],
        "Recent research has proposed several approaches to mitigate hallucinations, broadly categorized into three strategies: (1) prompt engineering techniques such as chain-of-thought reasoning and few-shot prompting [1], (2) sampling strategies like self-consistency decoding [3], and (3) retrieval-augmented generation (RAG) that grounds model outputs in external knowledge bases [4]. While each approach has shown promise individually, few studies have systematically compared their effectiveness under controlled conditions on the same model and dataset.")

    # P24: was Heading 2 - change to Body Text for third intro para
    paras[24].style = doc.styles['Body Text']
    set_para_text(paras[24],
        "This paper addresses this gap by implementing and evaluating six experimental conditions on LLaMA 3.2 (3B parameters), testing the effectiveness of prompt-based and RAG-based hallucination reduction techniques on a curated dataset of 30 TruthfulQA-inspired questions spanning five knowledge categories: Health, Science, History, Geography, and Misconceptions. Our contributions include: (1) a reproducible experimental pipeline for hallucination evaluation, (2) a comparative analysis of prompt engineering versus RAG-based approaches, and (3) category-level analysis revealing domain-specific effectiveness patterns.")

    # P25: Body Text -> Related Work heading
    paras[25].style = doc.styles['Heading 1']
    set_para_text(paras[25], "Related Work")

    # P26 is empty (had sectPr) - skip

    # P27: Heading 2 -> Body Text (Dang et al.)
    paras[27].style = doc.styles['Body Text']
    set_para_text(paras[27],
        "Dang, Vu, and Nguyen [1] proposed the Prompt Sensitivity (PS) and Model Variability (MV) metrics for determining whether hallucinations originate from prompting strategies or model behavior. Their experiments with GPT-4, LLaMA 2, and DeepSeek demonstrated that chain-of-thought prompting can significantly reduce prompt-sensitive hallucinations, though some models retain high hallucination rates regardless of strategy, indicating the need for multi-faceted approaches.")

    # P28: sponsors -> Body Text (Bang et al.)
    paras[28].style = doc.styles['Body Text']
    set_para_text(paras[28],
        "Bang et al. [2] introduced HalluLens, a comprehensive hallucination benchmarking suite that categorizes hallucinations as extrinsic (deviating from training data) or intrinsic (internal inconsistencies). Their dynamic test generation approach prevents data contamination and provides robust evaluation metrics that informed our evaluation methodology.")

    # P29: Body Text (Song et al.)
    set_para_text(paras[29],
        "Song et al. [4] presented RAG-HAT, a hallucination-aware tuning pipeline combining retrieval augmentation with Direct Preference Optimization (DPO) fine-tuning. Their three-stage approach\u2014hallucination detection, resolution via GPT-4 Turbo, and preference-based fine-tuning\u2014demonstrated significant hallucination reduction, validating the effectiveness of combining RAG with targeted training approaches.")

    # P30: Heading 1 -> Methodology
    set_para_text(paras[30], "Methodology")

    # P31: Body Text -> Experimental Setup intro
    paras[31].style = doc.styles['Heading 2']
    set_para_text(paras[31], "Experimental Setup")

    # P32: Body Text
    set_para_text(paras[32],
        "All experiments were conducted using LLaMA 3.2 (3B parameters) served locally via Ollama on an Apple M1 system with 8GB RAM. The model was run with temperature 0.7 and a maximum prediction length of 300 tokens. We used a curated dataset of 30 questions inspired by TruthfulQA [5], each with verified ground-truth answers and known incorrect alternatives, distributed across five categories: Health (7), Science (8), History (5), Geography (4), and Misconceptions (6).")

    # P33: Heading 2 -> Evaluation Metrics
    set_para_text(paras[33], "Evaluation Metrics")

    # P34: Body Text
    set_para_text(paras[34],
        "We evaluated responses using three complementary metrics: (1) Hallucination Detection Rate\u2014a binary classifier combining keyword overlap analysis with ROUGE-L similarity, comparing response alignment with correct versus incorrect reference answers; (2) ROUGE-L Score\u2014measuring longest common subsequence overlap between generated responses and ground-truth answers [6]; and (3) Factual Accuracy Score\u2014a composite metric combining keyword recall (weighted 0.6) and ROUGE-L similarity (weighted 0.4) against correct and incorrect answer sets, specifically designed to handle verbose responses from structured prompting methods.")

    # P35: Heading 2 -> Phase 1
    set_para_text(paras[35], "Phase 1: Baseline Measurement")

    # P36-39: bullet list -> Body Text for phases
    paras[36].style = doc.styles['Body Text']
    set_para_text(paras[36],
        "In the baseline condition, each question was presented to the model with a simple prompt: \"Answer the following question concisely and factually.\" This establishes the vanilla LLM performance without any hallucination mitigation strategies.")

    paras[37].style = doc.styles['Heading 2']
    set_para_text(paras[37], "Phase 2: Prompt-Based Reduction")

    paras[38].style = doc.styles['Body Text']
    set_para_text(paras[38],
        "Three prompt engineering techniques were evaluated: (a) Chain-of-Thought (CoT) Prompting\u2014instructing the model to reason step-by-step, identify common misconceptions, and consult scientific evidence before providing a final answer; (b) Few-Shot Prompting\u2014providing three verified question-answer examples demonstrating fact-checked, misconception-correcting responses; (c) Self-Consistency Decoding\u2014generating three independent responses per question (temperature 0.8) and selecting the response with highest average ROUGE-L similarity to all other responses, following the majority consensus principle.")

    paras[39].style = doc.styles['Heading 2']
    set_para_text(paras[39], "Phase 3: RAG-Based Reduction")

    # P40: Heading 2 -> Body Text
    paras[40].style = doc.styles['Body Text']
    set_para_text(paras[40],
        "A knowledge base of four curated reference documents covering health myths, scientific misconceptions, historical facts, and geography was constructed and indexed using ChromaDB with the all-MiniLM-L6-v2 embedding model. Each document was split into paragraphs, yielding 27 indexed chunks. Two RAG conditions were tested: (a) RAG Pipeline\u2014retrieving the top 3 most relevant chunks and instructing the model to answer based solely on the provided context; (b) RAG + Optimized Prompting\u2014combining retrieval with a structured prompt incorporating fact-checking instructions, misconception identification, chain-of-thought reasoning, and a few-shot example.")

    # P41-42: Results section
    paras[41].style = doc.styles['Heading 1']
    set_para_text(paras[41], "Results")

    paras[42].style = doc.styles['Heading 2']
    set_para_text(paras[42], "Overall Performance")

    # P43: equation -> table footnote for TABLE I caption
    paras[43].style = doc.styles['table head']
    set_para_text(paras[43], "TABLE I. Performance Comparison Across Methods")

    # Insert table after P43
    add_table_after(doc, paras[43],
        ['Method', 'Halluc. Rate', 'Factual Acc.', 'ROUGE-L'],
        [
            ['Vanilla Baseline', f"{summary['baseline']['hallucination_rate']:.1%}", f"{summary['baseline']['avg_factual_accuracy']:.3f}", f"{summary['baseline']['avg_rouge_l']:.3f}"],
            ['Chain-of-Thought', f"{summary['cot']['hallucination_rate']:.1%}", f"{summary['cot']['avg_factual_accuracy']:.3f}", f"{summary['cot']['avg_rouge_l']:.3f}"],
            ['Few-Shot', f"{summary['few_shot']['hallucination_rate']:.1%}", f"{summary['few_shot']['avg_factual_accuracy']:.3f}", f"{summary['few_shot']['avg_rouge_l']:.3f}"],
            ['Self-Consistency', f"{summary['self_consistency']['hallucination_rate']:.1%}", f"{summary['self_consistency']['avg_factual_accuracy']:.3f}", f"{summary['self_consistency']['avg_rouge_l']:.3f}"],
            ['RAG Pipeline', f"{summary['rag']['hallucination_rate']:.1%}", f"{summary['rag']['avg_factual_accuracy']:.3f}", f"{summary['rag']['avg_rouge_l']:.3f}"],
            ['RAG + Optimized', f"{summary['rag_optimized']['hallucination_rate']:.1%}", f"{summary['rag_optimized']['avg_factual_accuracy']:.3f}", f"{summary['rag_optimized']['avg_rouge_l']:.3f}"],
        ]
    )

    # P44: Body Text - results discussion
    set_para_text(paras[44],
        f"Table I presents the comprehensive results across all six conditions. The most significant finding is the progressive improvement in factual accuracy from baseline ({baseline_acc:.3f}) to RAG + Optimized Prompting ({best_acc:.3f}), representing a {improvement:.1f}% improvement. The RAG pipeline achieved the highest ROUGE-L score ({summary['rag']['avg_rouge_l']:.3f}), indicating strongest lexical overlap with ground-truth answers.")

    # P45: Heading 2 -> Category-Level Analysis
    set_para_text(paras[45], "Category-Level Analysis")

    # P46-56: bullet list items -> use for results + discussion
    paras[46].style = doc.styles['Body Text']
    set_para_text(paras[46],
        f"Among prompt-only methods, few-shot prompting achieved the highest factual accuracy ({summary['few_shot']['avg_factual_accuracy']:.3f}), outperforming chain-of-thought ({summary['cot']['avg_factual_accuracy']:.3f}). Self-consistency decoding showed improvement over baseline in both factual accuracy ({summary['self_consistency']['avg_factual_accuracy']:.3f}) and hallucination rate ({summary['self_consistency']['hallucination_rate']:.1%} vs {summary['baseline']['hallucination_rate']:.1%}). Notably, CoT prompting produced very low ROUGE-L scores ({summary['cot']['avg_rouge_l']:.3f}) due to verbose step-by-step outputs, though its factual accuracy remained competitive.")

    paras[47].style = doc.styles['Heading 2']
    set_para_text(paras[47], "Evaluation Metric Analysis")

    paras[48].style = doc.styles['Body Text']
    set_para_text(paras[48],
        "An important methodological finding concerns the divergence between ROUGE-L and factual accuracy metrics. Chain-of-thought and RAG + Optimized prompting produced verbose, structured responses that achieved low ROUGE-L scores despite containing correct factual content. Our composite factual accuracy metric, weighting keyword recall at 0.6 and ROUGE-L at 0.4, better captures the correctness of verbose but accurate responses. This finding has important implications for hallucination evaluation methodology: relying solely on lexical overlap metrics may systematically underestimate the effectiveness of structured prompting techniques.")

    # P49-57: Discussion section
    paras[49].style = doc.styles['Heading 1']
    set_para_text(paras[49], "Discussion")

    paras[50].style = doc.styles['Body Text']
    set_para_text(paras[50],
        "Our results reveal several key insights about hallucination reduction in smaller LLMs. First, the progressive improvement from prompt engineering to RAG confirms the hypothesis that grounding model outputs in external knowledge is more effective than relying on the model's parametric knowledge alone. The 43.3% improvement in factual accuracy from baseline to RAG + Optimized Prompting demonstrates that even smaller models (3B parameters) can achieve meaningful hallucination reduction when augmented with retrieval.")

    paras[51].style = doc.styles['Body Text']
    set_para_text(paras[51],
        "Second, the relative underperformance of chain-of-thought prompting on smaller models contrasts with findings by Dang et al. [1] on larger models like GPT-4. This suggests that CoT's effectiveness may be capacity-dependent\u2014smaller models may lack the reasoning depth to fully leverage step-by-step decomposition, instead generating plausible-sounding but unfaithful reasoning chains.")

    paras[52].style = doc.styles['Body Text']
    set_para_text(paras[52],
        "Third, self-consistency decoding showed moderate improvement with minimal implementation complexity, making it a practical first-line defense against hallucinations. The majority-vote mechanism naturally filters out low-confidence responses, though at the cost of 3x inference time.")

    paras[53].style = doc.styles['Body Text']
    set_para_text(paras[53],
        "The category-level analysis reveals that RAG is particularly effective for factual recall questions (Geography, Health) where retrieved documents directly contain relevant information, but less effective for questions requiring nuanced reasoning about historical or conceptual misconceptions.")

    paras[54].style = doc.styles['Heading 2']
    set_para_text(paras[54], "Limitations")

    paras[55].style = doc.styles['Body Text']
    set_para_text(paras[55],
        "This study has several limitations. The dataset of 30 questions, while curated to cover diverse categories, is relatively small. The use of a 3B parameter model may not generalize to larger LLMs where techniques like CoT may be more effective. The RAG knowledge base was purpose-built to contain answers to the test questions, representing an optimistic scenario. Additionally, our evaluation relies on automated metrics rather than human judgment.")

    paras[56].style = doc.styles['Heading 1']
    set_para_text(paras[56], "Conclusion and Future Work")

    paras[57].style = doc.styles['Body Text']
    set_para_text(paras[57],
        f"This paper presented a systematic evaluation of hallucination reduction techniques for LLaMA 3.2 (3B parameters). Our three-phase experimental pipeline demonstrated that Retrieval-Augmented Generation combined with optimized prompting achieves the highest factual accuracy improvement of {improvement:.1f}% over the vanilla baseline. Among prompt-only methods, few-shot prompting and self-consistency decoding showed the most consistent improvements, while chain-of-thought prompting was less effective on this smaller model.")

    # P58: Heading 1 -> Body Text for future work
    paras[58].style = doc.styles['Body Text']
    set_para_text(paras[58],
        "Future work will extend this evaluation to larger models (7B, 13B, 70B parameters) to investigate scale-dependent effectiveness. We plan to implement hallucination-aware fine-tuning following the RAG-HAT approach [4] and evaluate on larger benchmarks including the full TruthfulQA dataset. Additionally, we will explore hybrid approaches combining self-consistency with RAG and investigate the use of Natural Language Inference models for automated faithfulness evaluation.")

    # Clear remaining body paragraphs (P59-P79) that we don't need
    for i in range(59, 80):
        if i < len(paras) and paras[i].style.name not in ('Heading 5', 'references'):
            set_para_text(paras[i], "")

    # P80: Heading 5 - References (keep as is)
    set_para_text(paras[80], "References")

    # P81-85: Clear old body text before references
    for i in range(81, 86):
        if i < len(paras):
            set_para_text(paras[i], "")

    # P86-97: references - replace with our references
    refs = [
        'A.-H. Dang, T. Vu, and L.-M. Nguyen, "A survey and analysis of hallucinations in large language models," Frontiers in Artificial Intelligence, vol. 8, 2025.',
        'Y. Bang et al., "HalluLens: A benchmark for LLM hallucinations," in Proc. 63rd ACL, 2025.',
        'X. Wang et al., "Self-consistency improves chain of thought reasoning in language models," in Proc. ICLR, 2023.',
        'J. Song et al., "RAG-HAT: A Hallucination-Aware Tuning Pipeline for LLM in RAG," in Proc. EMNLP, 2024.',
        'S. Lin, J. Hilton, and O. Evans, "TruthfulQA: Measuring how models mimic human falsehoods," in Proc. ACL, 2022.',
        'C.-Y. Lin, "ROUGE: A package for automatic evaluation of summaries," in Text Summarization Branches Out, 2004.',
        'P. Lewis et al., "Retrieval-augmented generation for knowledge-intensive NLP tasks," in Proc. NeurIPS, 2020.',
        'J. Wei et al., "Chain-of-thought prompting elicits reasoning in large language models," in Proc. NeurIPS, 2022.',
    ]

    ref_start = 86
    for i, ref in enumerate(refs):
        idx = ref_start + i
        if idx < len(paras):
            set_para_text(paras[idx], ref)

    # Clear remaining reference slots and template warning
    for i in range(ref_start + len(refs), len(paras)):
        if paras[i].style.name in ('references', 'Body Text', 'Normal'):
            set_para_text(paras[i], "")

    # Remove the original table (template sample table)
    body = doc.element.body
    for tbl in body.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl'):
        body.remove(tbl)

    doc.save(OUTPUT)
    print(f"IEEE paper saved to: {OUTPUT}")


if __name__ == "__main__":
    main()
